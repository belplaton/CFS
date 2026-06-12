"""
Preview Service.

This service generates text-oriented previews for formats browsers do
not render well on their own, while browser-native previews (images /
PDFs) can still be fetched directly from the file service.
"""
from __future__ import annotations

from io import BytesIO
import json
import uuid

import httpx
import structlog
from fastapi import FastAPI, Header, HTTPException, Request, status
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from openpyxl import load_workbook
from pydantic import BaseModel

from src.config import settings


logger = structlog.get_logger()

TEXT_MIME_TYPES = {
    "text/plain",
    "text/csv",
    "application/json",
}

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


class TextPreviewResponse(BaseModel):
    kind: str = "text"
    content: str
    truncated: bool = False


app = FastAPI(
    title="Cloud Storage Preview Service",
    description="File preview generation service for Cloud File Storage",
    version="1.0.0",
    docs_url="/docs/preview",
    redoc_url="/redoc/preview",
    openapi_url="/openapi/preview.json",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def add_request_id(request: Request, call_next):
    request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())
    structlog.contextvars.bind_contextvars(request_id=request_id)
    response = await call_next(request)
    response.headers["X-Request-ID"] = request_id
    return response


@app.get("/health")
async def health_check():
    file_service_ok = True
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{settings.file_service_url.rstrip('/')}/health")
            file_service_ok = resp.status_code == 200
    except Exception:
        file_service_ok = False

    healthy = file_service_ok
    return {
        "status": "healthy" if healthy else "degraded",
        "service": "preview",
        "file_service": "healthy" if file_service_ok else "unreachable",
    }


@app.get("/api/preview/")
async def root():
    return {
        "message": "Preview Service is running",
        "version": "1.0.0",
        "generated_previews_enabled": True,
        "supported_text_previews": [
            "txt",
            "csv",
            "json",
            "docx",
            "xlsx",
        ],
        "note": "Images and PDFs can still use direct file download preview.",
    }


def _require_auth_header(authorization: str | None) -> str:
    if not authorization:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing Authorization header",
        )
    return authorization


async def _fetch_file_bytes(file_id: str, authorization: str) -> tuple[bytes, str]:
    url = f"{settings.file_service_url.rstrip('/')}/api/files/{file_id}/download"
    headers = {"Authorization": authorization}
    if settings.service_api_key:
        headers["X-API-Key"] = settings.service_api_key

    async with httpx.AsyncClient(timeout=30.0) as client:
        try:
            response = await client.get(url, headers=headers)
        except httpx.TimeoutException:
            raise HTTPException(
                status_code=status.HTTP_504_GATEWAY_TIMEOUT,
                detail="File service request timed out",
            )
        except httpx.RequestError:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail="Preview service cannot reach file service",
            )

        if response.status_code >= 400:
            detail = response.text or "Unable to fetch source file"
            raise HTTPException(status_code=response.status_code, detail=detail)

        content = response.content
        if len(content) > settings.preview_max_size:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File exceeds maximum preview size of {settings.preview_max_size} bytes",
            )

        mime_type = response.headers.get("content-type", "").split(";")[0].strip()
        return content, mime_type


def _limit_text(value: str, max_chars: int = 40000) -> tuple[str, bool]:
    if len(value) <= max_chars:
        return value, False
    return value[:max_chars], True


def _preview_text_bytes(content: bytes) -> TextPreviewResponse:
    text = content.decode("utf-8", errors="replace")
    limited, truncated = _limit_text(text)
    return TextPreviewResponse(content=limited, truncated=truncated)


def _preview_json_bytes(content: bytes) -> TextPreviewResponse:
    parsed = json.loads(content.decode("utf-8", errors="replace"))
    pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
    limited, truncated = _limit_text(pretty)
    return TextPreviewResponse(content=limited, truncated=truncated)


def _preview_docx_bytes(content: bytes) -> TextPreviewResponse:
    document = Document(BytesIO(content))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    preview_text = "\n\n".join(lines).strip() or "Document contains no extractable text."
    limited, truncated = _limit_text(preview_text)
    return TextPreviewResponse(content=limited, truncated=truncated)


def _preview_xlsx_bytes(content: bytes) -> TextPreviewResponse:
    workbook = load_workbook(filename=BytesIO(content), read_only=True, data_only=True)
    sheet = workbook.worksheets[0]
    rows: list[str] = [f"Sheet: {sheet.title}"]

    for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
        if index > 100:
            rows.append("...")
            break
        formatted = ["" if cell is None else str(cell) for cell in row]
        rows.append("\t".join(formatted).rstrip())

    preview_text = "\n".join(rows).strip() or "Spreadsheet contains no previewable cells."
    limited, truncated = _limit_text(preview_text)
    return TextPreviewResponse(content=limited, truncated=truncated)


@app.get("/api/preview/{file_id}", response_model=TextPreviewResponse)
async def get_preview(file_id: str, authorization: str | None = Header(default=None)):
    auth_header = _require_auth_header(authorization)

    logger.info("preview_request", file_id=file_id)

    content, mime_type = await _fetch_file_bytes(file_id, auth_header)

    if mime_type == "application/json":
        return _preview_json_bytes(content)

    if mime_type in TEXT_MIME_TYPES or mime_type.startswith("text/"):
        return _preview_text_bytes(content)

    if mime_type == DOCX_MIME_TYPE:
        return _preview_docx_bytes(content)

    if mime_type == XLSX_MIME_TYPE:
        return _preview_xlsx_bytes(content)

    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="Preview is not available for this file type",
    )


@app.get("/api/preview/{file_id}/thumbnail")
async def get_thumbnail(
    file_id: str,
    authorization: str | None = Header(default=None),
):
    _require_auth_header(authorization)
    raise HTTPException(
        status_code=status.HTTP_501_NOT_IMPLEMENTED,
        detail="Thumbnails are not enabled yet",
    )


@app.post("/api/preview/{file_id}/generate")
async def generate_preview(
    file_id: str,
    authorization: str | None = Header(default=None),
):
    _require_auth_header(authorization)
    raise HTTPException(
        status_code=status.HTTP_501_NOT_IMPLEMENTED,
        detail="Background preview generation is not enabled yet",
    )


@app.delete("/api/preview/{file_id}")
async def delete_preview(
    file_id: str,
    authorization: str | None = Header(default=None),
):
    _require_auth_header(authorization)
    raise HTTPException(
        status_code=status.HTTP_501_NOT_IMPLEMENTED,
        detail="Stored previews are not enabled yet",
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
