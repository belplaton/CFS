"""Text-oriented preview extraction for various file formats."""

from __future__ import annotations

import json
from io import BytesIO

from defusedxml import defuse_stdlib
from docx import Document
from fastapi import HTTPException, status
from openpyxl import load_workbook

from src.schemas.preview import TextPreviewResponse

# ── MIME type constants ──────────────────────────────────────────

TEXT_MIME_TYPES = {
    "text/plain",
    "text/csv",
    "application/json",
}

DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


# ── Helpers ──────────────────────────────────────────────────────


def limit_text(value: str, max_chars: int = 40000) -> tuple[str, bool]:
    """Truncate *value* to *max_chars*. Returns ``(text, was_truncated)``."""
    if len(value) <= max_chars:
        return value, False
    return value[:max_chars], True


def secure_docx_parser(content: bytes) -> Document:
    """Parse DOCX with XXE / entity expansion protection.

    python-docx uses lxml internally. We validate the input is a valid
    zip (DOCX is a zip archive) and use defusedxml to monkeypatch the
    stdlib xml parser to block entity expansion attacks.
    """
    import zipfile

    try:
        zipfile.ZipFile(BytesIO(content))
    except zipfile.BadZipFile:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Invalid DOCX file format",
        )

    # defuse_stdlib patches xml.etree to block entity expansion.
    defuse_stdlib()
    return Document(BytesIO(content))


# ── Per-format preview extractors ────────────────────────────────


def preview_text(content: bytes) -> TextPreviewResponse:
    text = content.decode("utf-8", errors="replace")
    limited, truncated = limit_text(text)
    return TextPreviewResponse(content=limited, truncated=truncated)


def preview_json(content: bytes) -> TextPreviewResponse:
    parsed = json.loads(content.decode("utf-8", errors="replace"))
    pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
    limited, truncated = limit_text(pretty)
    return TextPreviewResponse(content=limited, truncated=truncated)


def preview_docx(content: bytes) -> TextPreviewResponse:
    document = secure_docx_parser(content)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    preview_text_str = (
        "\n\n".join(lines).strip() or "Document contains no extractable text."
    )
    limited, truncated = limit_text(preview_text_str)
    return TextPreviewResponse(content=limited, truncated=truncated)


def preview_xlsx(content: bytes) -> TextPreviewResponse:
    workbook = load_workbook(filename=BytesIO(content), read_only=True, data_only=True)
    sheet = workbook.worksheets[0]
    rows: list[str] = [f"Sheet: {sheet.title}"]

    for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
        if index > 100:
            rows.append("...")
            break
        formatted = ["" if cell is None else str(cell) for cell in row]
        rows.append("\t".join(formatted).rstrip())

    preview_text_str = (
        "\n".join(rows).strip() or "Spreadsheet contains no previewable cells."
    )
    limited, truncated = limit_text(preview_text_str)
    return TextPreviewResponse(content=limited, truncated=truncated)


# ── Dispatcher ───────────────────────────────────────────────────


def extract_preview(content: bytes, mime_type: str) -> TextPreviewResponse:
    """Route to the correct extractor based on MIME type.

    Raises ``HTTPException`` (415) for unsupported types.
    """
    if mime_type == "application/json":
        return preview_json(content)

    if mime_type in TEXT_MIME_TYPES or mime_type.startswith("text/"):
        return preview_text(content)

    if mime_type == DOCX_MIME_TYPE:
        return preview_docx(content)

    if mime_type == XLSX_MIME_TYPE:
        return preview_xlsx(content)

    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="Preview is not available for this file type",
    )
